from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_PATH = Path("docs/Huong_dan_su_dung_Phong_ban_Luong_Thu_mua.docx")


TITLE = "HƯỚNG DẪN SỬ DỤNG PHÂN HỆ PHÒNG BAN, LƯƠNG VÀ THU MUA"
SUBTITLE = (
    "Tài liệu này mô tả chức năng, công dụng, trạng thái, vai trò và cách vận hành "
    "của 3 phân hệ Phòng ban, Lương và Thu mua theo hệ thống hiện tại."
)


PAYROLL_PERIOD_STATUS = [
    ("Nháp", "draft", "Kỳ lương đang soạn, còn được tính lại và sửa số."),
    ("Đã chốt", "locked", "Đã khóa số liệu, không cho sửa trực tiếp nữa."),
    ("Đã chi", "paid", "Đã xác nhận chi trả, khóa cứng hơn kỳ đã chốt."),
]

ADVANCE_STATUS = [
    ("Chờ duyệt", "pending", "Phiếu mới tạo, đang chờ người có quyền duyệt."),
    ("Đã duyệt", "approved", "Phiếu hợp lệ, số tiền sẽ được tính để trừ vào lương."),
    ("Từ chối", "rejected", "Phiếu không được duyệt, không đi vào bảng lương."),
    ("Đã hủy", "cancelled", "Phiếu bị hủy, ngừng hiệu lực."),
]

SCOPE_ROWS = [
    ("Của tôi", "Chỉ đụng tới dữ liệu của chính người dùng."),
    ("Cả phòng", "Đụng tới dữ liệu thuộc phòng hoặc tổ của mình và các tổ con."),
    ("Tất cả", "Đụng tới dữ liệu toàn công ty."),
]

ROLE_ROWS = [
    ("Admin", "Xem và thao tác toàn bộ, gồm cơ cấu, quyền, lương, tạm ứng, thu mua, duyệt và xuất file."),
    ("HCNS / C&B", "Quản lý hồ sơ, khai lương, tính lương, soát bảng lương, cấu hình lương."),
    ("Kế toán lương", "Duyệt tạm ứng, theo dõi chi trả, xác nhận đã chi, xuất dữ liệu chuyển khoản."),
    ("Thu mua", "Nhận yêu cầu mua hàng từ các phòng ban, lập phiếu mua hàng, theo dõi đã mua và đã nhận."),
    ("Người duyệt / kế toán mua hàng", "Duyệt hoặc từ chối phiếu mua hàng, lập chứng từ chi ở phân hệ kế toán khi cần."),
    ("Trưởng phòng / trưởng bộ phận", "Quản lý nhân sự trong phạm vi được cấp, có thể được giao thêm quyền vận hành hoặc duyệt."),
    ("Nhân viên", "Xem phiếu lương của mình, theo dõi và tạo đề nghị tạm ứng của mình."),
]

ROOM_SUMMARY_ROWS = [
    ("Tổng quan", "Nhìn nhanh toàn bộ cơ cấu, số lượng phòng, nhân sự, tình trạng trưởng phòng."),
    ("Nhân sự", "Xem người trong phòng, thêm nhân viên, gán vai trò hàng loạt, điều chuyển phòng."),
    ("Vai trò & Quyền", "Tạo vai trò của từng phòng và gán ma trận phân quyền chi tiết."),
]

PAYROLL_TABS = [
    ("Bảng lương tháng", "Tạo bảng lương, soát số, chốt kỳ, xuất file, in phiếu."),
    ("Lương nhân viên", "Khai mức lương và các khoản cố định của từng nhân viên."),
    ("Lương khoán", "Khu vực dành cho luồng lương khoán khi doanh nghiệp áp dụng."),
    ("Tạm ứng", "Quản lý phiếu tạm ứng và phiếu lương đợt 1."),
    ("Cấu hình lương", "Quản lý tham số chung, cấu hình theo bộ phận, bảo hiểm, thuế, phụ cấp."),
    ("Phiếu lương của tôi", "Nhân viên tự xem phiếu lương cá nhân."),
    ("Tạm ứng của tôi", "Nhân viên tự tạo và theo dõi đề nghị tạm ứng của chính mình."),
]

PURCHASE_AREAS = [
    ("Yêu cầu mua hàng", "Nơi các phòng ban tạo YCMH, theo dõi yêu cầu của chính mình và chỉnh sửa khi chưa được Thu mua lập PMH."),
    ("Mua hàng", "Nơi bộ phận Thu mua nhận YCMH, tạo PMH từ từng yêu cầu, theo dõi trạng thái đã mua và đã nhận."),
    ("Nhà cung cấp", "Quản lý danh sách nhà cung cấp, thông tin liên hệ và mặt hàng - đơn giá mà từng nhà cung cấp đang bán."),
]

PURCHASE_STATUS = [
    ("Nháp", "draft", "Phiếu mua hàng mới tạo, còn chỉnh sửa được."),
    ("Chờ duyệt", "pending_approval", "Phiếu đã gửi lên người có quyền duyệt, chưa được quyết định."),
    ("Đã duyệt", "approved", "Phiếu hợp lệ để đi tiếp sang bước mua và/hoặc kế toán."),
    ("Từ chối", "rejected", "Phiếu không được duyệt, có thể cần sửa lại theo lý do phản hồi."),
    ("Đã mua", "purchased", "Thu mua xác nhận đã đặt hoặc đã mua xong với nhà cung cấp."),
    ("Đã nhận", "received", "Hàng đã nhận thực tế, có thể là đủ hoặc có số nhận thực tế khác số đặt."),
    ("Đã hủy", "cancelled", "Phiếu dừng hẳn, không còn hiệu lực."),
]

THU_MUA_PERMISSION_ROWS = [
    (
        "Xem",
        "Menu Thu mua xuất hiện; mở được màn Yêu cầu mua hàng, Mua hàng và danh sách liên quan trong phạm vi được cấp.",
        "Theo dõi YCMH từ các phòng ban, xem PMH đã lập, xem trạng thái duyệt, đã mua, đã nhận.",
        "Chỉ có Xem thì chưa tạo được yêu cầu và chưa lập được phiếu.",
    ),
    (
        "Thao tác",
        "Hiện nút tạo YCMH, tạo PMH, sửa phiếu nháp hoặc phiếu bị từ chối, gửi duyệt, đánh dấu đã mua, đánh dấu đã nhận ở nơi phù hợp.",
        "Vận hành luồng thu mua hằng ngày từ lúc nhận yêu cầu tới lúc xác nhận hàng về.",
        "Không tự bao gồm quyền duyệt hoặc hủy PMH nếu chưa bật quyền chi tiết.",
    ),
    (
        "Duyệt / từ chối PMH",
        "Hiện các nút Duyệt và Từ chối trên các phiếu đang chờ duyệt.",
        "Quyết định phiếu mua hàng nào được đi tiếp sang bước chi tiền hoặc mua hàng thực tế.",
        "Đây là quyền quyết định nghiệp vụ, nên tách khỏi người chỉ lập phiếu.",
    ),
    (
        "Hủy PMH",
        "Hiện thao tác hủy phiếu mua hàng còn hợp lệ.",
        "Dùng khi nhu cầu mua dừng hẳn hoặc phiếu lập sai và không muốn tiếp tục dùng.",
        "Nên yêu cầu nhập lý do để truy vết vì hủy sẽ dừng cả luồng phía sau.",
    ),
]

PERMISSION_MATRIX_OVERVIEW = [
    (
        "Xem",
        "Khi bật thì module xuất hiện trên menu hoặc cho phép mở màn hình tương ứng, nhìn thấy danh sách, bộ lọc, tab và chi tiết dữ liệu trong phạm vi được cấp.",
    ),
    (
        "Thao tác",
        "Khi bật thì người dùng được thêm, sửa, xóa hoặc thực hiện các thao tác nghiệp vụ chính của module đó. Nếu chỉ có Xem mà không có Thao tác thì chỉ đọc được dữ liệu.",
    ),
    (
        "Phạm vi",
        "Quy định người dùng được đụng tới dữ liệu của ai: chỉ của chính mình, của cả phòng mình, hay toàn công ty.",
    ),
]

PHONG_BAN_PERMISSION_ROWS = [
    (
        "Xem",
        "Menu Phòng ban xuất hiện; mở được màn Tổng quan, cây tổ chức, danh sách phòng, tab Nhân sự và tab Vai trò & Quyền của đơn vị nằm trong phạm vi.",
        "Theo dõi cơ cấu tổ chức, số lượng nhân sự, trạng thái trưởng phòng, xem các phòng con và người thuộc từng phòng.",
        "Tắt quyền này thì các quyền khác của module Phòng ban cũng không còn giá trị sử dụng.",
    ),
    (
        "Thao tác",
        "Hiện các nút thêm, sửa, xóa phòng ban ở nơi người dùng được phép thao tác.",
        "Tạo phòng ban mới, chỉnh tên/mã/ghi chú, xóa đơn vị không còn dùng.",
        "Không tự bao gồm quyền đổi cấp trên hay đặt trưởng phòng nếu chưa bật quyền chi tiết tương ứng.",
    ),
    (
        "Đặt trưởng phòng",
        "Hiện phần chọn hoặc thay đổi người đứng đầu đơn vị.",
        "Gán trưởng phòng, thay trưởng phòng, bỏ gán người đứng đầu cũ.",
        "Nên cấp cho Admin hoặc HCNS phụ trách cơ cấu tổ chức.",
    ),
    (
        "Đổi cấp trên",
        "Hiện thao tác đổi phòng cha trong cây tổ chức.",
        "Di chuyển một phòng hoặc tổ sang nhánh khác trong sơ đồ công ty.",
        "Nên dùng cẩn thận vì sẽ ảnh hưởng cách nhìn dữ liệu theo cây phòng ban.",
    ),
]

VAI_TRO_PERMISSION_ROWS = [
    (
        "Xem",
        "Mở được danh sách vai trò của phòng, xem ma trận quyền đã gán cho từng vai trò.",
        "Dùng khi cần kiểm tra một vai trò hiện có quyền gì mà chưa cần sửa.",
        "Không có quyền này thì người dùng không thấy cấu trúc vai trò của phòng.",
    ),
    (
        "Thao tác",
        "Hiện nút thêm, sửa tên, xóa vai trò.",
        "Tạo chỗ ngồi quyền mới cho từng phòng ban, ví dụ Trưởng nhóm, Kế toán lương, C&B.",
        "Thao tác này chưa có nghĩa là sửa được ma trận quyền chi tiết bên trong.",
    ),
    (
        "Sửa ma trận phân quyền",
        "Cho phép bật/tắt các ô quyền, đổi phạm vi của từng module trong vai trò.",
        "Đây là quyền mạnh nhất của phần Vai trò vì nó quyết định người khác nhìn thấy và làm được gì trong hệ thống.",
        "Nên chỉ cấp cho Admin hoặc người được ủy quyền quản trị phân quyền.",
    ),
]

NGUOI_DUNG_PERMISSION_ROWS = [
    (
        "Xem",
        "Thấy tab Tài khoản & Quyền trong hồ sơ nhân sự và các thông tin đăng nhập liên quan.",
        "Kiểm tra nhân viên đã có tài khoản hay chưa, đang giữ vai trò nào.",
        "Không bao gồm đổi mật khẩu hay đổi vai trò nếu chưa bật quyền chi tiết tương ứng.",
    ),
    (
        "Thao tác",
        "Hiện các nút tạo hoặc cập nhật tài khoản ở nơi hệ thống cho phép.",
        "Tạo tài khoản đăng nhập, chỉnh thông tin tài khoản cơ bản.",
        "Các thao tác nhạy cảm vẫn tách riêng bằng quyền chi tiết bên dưới.",
    ),
    (
        "Gán vai trò",
        "Hiện ô chọn vai trò cho tài khoản nhân viên.",
        "Đổi người dùng sang vai trò đúng của phòng ban đó.",
        "Vai trò phải thuộc đúng phòng của nhân viên, hệ thống không cho gán chéo phòng bừa bãi.",
    ),
    (
        "Chuyển phòng ban",
        "Hiện thao tác chuyển nhân sự sang phòng khác trong các công cụ gán hàng loạt hoặc hồ sơ.",
        "Phục vụ điều chuyển tổ chức.",
        "Sau khi chuyển phòng cần rà lại vai trò vì vai trò cũ có thể không còn hợp lệ.",
    ),
    (
        "Đặt lại mật khẩu",
        "Hiện nút tạo mật khẩu mới cho tài khoản nhân viên.",
        "Dùng khi nhân viên quên mật khẩu hoặc bàn giao tài khoản.",
        "Nên đi kèm quy trình thông báo mật khẩu tạm và yêu cầu đổi sau lần đăng nhập đầu.",
    ),
    (
        "Khóa / mở khóa tài khoản",
        "Hiện công tắc khóa đăng nhập của tài khoản.",
        "Ngăn đăng nhập tạm thời khi nhân viên nghỉ việc, bị đình chỉ hoặc cần chặn truy cập.",
        "Khóa tài khoản không xóa hồ sơ nhân sự.",
    ),
    (
        "Thu hồi phiên",
        "Hiện thao tác ép đăng xuất khỏi các thiết bị đang còn đăng nhập.",
        "Dùng khi nghi ngờ lộ phiên hoặc cần buộc người dùng đăng nhập lại.",
        "Rất hữu ích sau khi đổi mật khẩu hoặc khi nhân sự nghỉ việc.",
    ),
]

NHAN_SU_PERMISSION_ROWS = [
    (
        "Xem",
        "Menu Hồ sơ nhân sự và các danh sách nhân viên xuất hiện; mở được hồ sơ, quá trình công tác, đính kèm, nhật ký trong phạm vi được cấp.",
        "Theo dõi thông tin nhân viên, tìm kiếm, lọc trạng thái, xem lịch sử làm việc.",
        "Chỉ có Xem thì chưa chỉnh hồ sơ, chưa đổi trạng thái và chưa sửa dữ liệu nhạy cảm.",
    ),
    (
        "Thao tác",
        "Hiện nút Thêm nhân viên, Sửa thông tin, xóa hoặc cập nhật các trường hồ sơ thông thường.",
        "Tạo hồ sơ nhân viên, sửa thông tin cá nhân, liên hệ, địa chỉ, ghi chú.",
        "Không tự bao gồm sửa lương, đổi trạng thái hay điều chuyển nếu chưa bật quyền chi tiết.",
    ),
    (
        "Xem lương & BHXH",
        "Hiện tab Lương & BHXH trong hồ sơ nhân sự; cho xem số sổ BHXH, MST cá nhân, người phụ thuộc, tài khoản ngân hàng, hoa hồng và các dữ liệu nhạy cảm liên quan.",
        "Dùng khi cần đọc dữ liệu nhạy cảm để kiểm tra hồ sơ hoặc đối chiếu lương.",
        "Nên cấp hạn chế vì đây là nhóm dữ liệu nhạy cảm.",
    ),
    (
        "Sửa lương & BHXH",
        "Bật được chế độ sửa trong tab Lương & BHXH và hiện bước Lương khi tạo nhân viên mới.",
        "Khai ban đầu hoặc cập nhật số sổ BHXH, MST, người phụ thuộc, tài khoản ngân hàng, lương khởi tạo và các trường nhạy cảm mà backend cho phép.",
        "Quyền này phải đi cùng quyền xem lương; khi tắt Xem lương thì quyền sửa cũng bị tắt theo.",
    ),
    (
        "Thao tác vòng đời",
        "Hiện các nút chuyển chính thức, cho nghỉ dài hạn, đi làm lại, đình chỉ, thôi việc, tuyển lại trong menu Thao tác hồ sơ.",
        "Quản lý trạng thái làm việc của nhân viên theo thực tế ngoài đời.",
        "Mỗi thao tác đều sinh mốc trong quá trình công tác để truy vết lịch sử.",
    ),
    (
        "Điều chuyển & nâng bậc",
        "Hiện các nút Điều chuyển phòng/tổ và Nâng bậc / Đổi chức danh trong menu Thao tác hồ sơ.",
        "Đổi phòng ban, đổi bậc tay nghề, đổi chức danh của nhân viên.",
        "Không dùng quyền này để sửa tiền lương; đổi tiền lương làm ở phân hệ Lương.",
    ),
    (
        "Duyệt yêu cầu cập nhật",
        "Hiện danh sách Yêu cầu cập nhật của nhân viên và nút duyệt/từ chối.",
        "Dùng để HCNS duyệt các yêu cầu nhân viên tự gửi từ Hồ sơ của tôi.",
        "Nếu người duyệt không có quyền sửa lương thì backend cũng không cho ghi đè nhóm field nhạy cảm.",
    ),
    (
        "Xuất Excel danh sách",
        "Hiện nút xuất danh sách nhân viên ra file.",
        "Phục vụ báo cáo nhanh, kiểm tra hoặc gửi nội bộ.",
        "Nên cấp cho người thật sự cần kéo dữ liệu ra ngoài hệ thống.",
    ),
    (
        "Chấm bù / sửa công",
        "Hiện các thao tác sửa công, chấm bù hoặc điều chỉnh chấm công ở các màn liên quan.",
        "Dùng khi dữ liệu chấm công thực tế cần hiệu chỉnh có kiểm soát.",
        "Quyền này ảnh hưởng gián tiếp đến lương nên nên cấp hạn chế.",
    ),
]

LUONG_PERMISSION_ROWS = [
    (
        "Xem",
        "Menu Lương xuất hiện; mở được Bảng lương tháng và Tạm ứng để xem dữ liệu trong phạm vi cho phép.",
        "Theo dõi bảng lương, xem các dòng lương, xem danh sách tạm ứng mà chưa cần sửa hay chốt.",
        "Chỉ có Xem thì chưa mở được tab Lương nhân viên hay thao tác sửa sâu.",
    ),
    (
        "Thao tác",
        "Hiện các tab Lương nhân viên, Lương khoán và các nút sửa dữ liệu lương ở nơi phù hợp.",
        "Tính lại bảng lương nháp, sửa dòng lương, khai hồ sơ lương từng người, cập nhật dữ liệu vận hành lương.",
        "Đây là quyền vận hành mạnh, thường cấp cho C&B hoặc người trực tiếp làm lương.",
    ),
    (
        "Xem cấu hình lương",
        "Hiện tab Cấu hình lương ngay cả khi người dùng không có quyền sửa vận hành lương.",
        "Cho xem các tham số, thuế, bảo hiểm, cấu hình theo bộ phận để kiểm tra hoặc đối chiếu.",
        "Chỉ xem, không tự lưu thay đổi nếu thiếu quyền Thao tác.",
    ),
    (
        "Duyệt tạm ứng",
        "Hiện các nút Duyệt, Từ chối, Hủy trên danh sách phiếu tạm ứng và hiện badge chờ duyệt.",
        "Quyết định phiếu tạm ứng nào hợp lệ để đưa vào phần khấu trừ lương.",
        "Tách riêng với quyền tạo phiếu để giữ kiểm soát 2 bước.",
    ),
    (
        "Chốt kỳ lương",
        "Hiện các nút Chốt, Mở lại, Đã chi, Hủy đã chi trên bảng lương tháng.",
        "Khóa số liệu lương của tháng và xác nhận trạng thái chi trả.",
        "Chỉ nên cấp cho người chịu trách nhiệm chốt cuối cùng vì đây là mốc nghiệp vụ quan trọng.",
    ),
    (
        "Xuất bảng lương / file chuyển khoản",
        "Hiện nút Xuất Excel và File chuyển khoản ngân hàng.",
        "Dùng để gửi kế toán hoặc tạo file chi lương qua ngân hàng.",
        "Nên cấp riêng thay vì mặc định ai xem lương cũng được xuất file.",
    ),
]


def set_cell_text(cell, text: str, *, bold: bool = False, center: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    set_run_font(run, size=Pt(10.5), color="222222")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, *, size: Pt | None = None, color: str | None = None, bold: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = size
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], widths_cm: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr = table.rows[0].cells
    for idx, text in enumerate(headers):
        hdr[idx].width = Cm(widths_cm[idx])
        set_cell_text(hdr[idx], text, bold=True)
        set_cell_shading(hdr[idx], "E8EEF5")

    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].width = Cm(widths_cm[idx])
            set_cell_text(cells[idx], text)

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)


def add_permission_table(
    doc: Document,
    title: str,
    intro: str,
    rows: list[tuple[str, str, str, str]],
) -> None:
    add_heading(doc, title, 3)
    add_para(doc, intro)
    add_table(
        doc,
        ["Quyền / nút", "Khi bật sẽ thấy gì", "Làm được gì / công dụng", "Lưu ý"],
        rows,
        [3.0, 5.0, 5.0, 3.5],
    )


def add_bullet(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if level:
        p.paragraph_format.left_indent = Cm(0.6 * level)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=Pt(11), color="222222")


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=Pt(11), color="222222")


def add_para(doc: Document, text: str, *, italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=Pt(11), color="222222")
    run.italic = italic


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6 if level == 1 else 4)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=Pt(16), color="2E74B5", bold=True)
    elif level == 2:
        set_run_font(run, size=Pt(13), color="2E74B5", bold=True)
    else:
        set_run_font(run, size=Pt(12), color="1F4D78", bold=True)


def add_footer_with_page_number(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.text = ""
    label = p.add_run("Trang ")
    set_run_font(label, size=Pt(9), color="666666")

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run = p.add_run()
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    txt = OxmlElement("w:t")
    txt.text = "1"
    run._r.append(txt)
    run._r.append(fld_end)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    add_footer_with_page_number(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)


def add_title_block(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(TITLE)
    set_run_font(run, size=Pt(20), color="1F3A5F", bold=True)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(12)
    r2 = p2.add_run(SUBTITLE)
    set_run_font(r2, size=Pt(11), color="4A5568")

    info = doc.add_table(rows=2, cols=2)
    info.style = "Table Grid"
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    info.autofit = False
    widths = [3.4, 12.0]
    rows = [
        ("Phạm vi", "Phân hệ Phòng ban, phân hệ Lương và phân hệ Thu mua"),
        ("Mục tiêu", "Giúp quản trị, HCNS, kế toán, Thu mua và nhân viên hiểu đúng cách dùng hệ thống hiện tại"),
    ]
    for ridx, row in enumerate(rows):
        for cidx, val in enumerate(row):
            info.rows[ridx].cells[cidx].width = Cm(widths[cidx])
            set_cell_text(info.rows[ridx].cells[cidx], val, bold=(cidx == 0))
            if cidx == 0:
                set_cell_shading(info.rows[ridx].cells[cidx], "F2F4F7")


def build_doc() -> Path:
    doc = Document()
    configure_styles(doc)
    add_title_block(doc)

    add_heading(doc, "1. Mục đích tài liệu", 1)
    add_para(
        doc,
        "Tài liệu này giải thích rõ chức năng, công dụng, trạng thái, vai trò và cách vận hành của "
        "ba phân hệ Phòng ban, Lương và Thu mua theo đúng luồng hiện có trong hệ thống."
    )
    add_para(
        doc,
        "Nội dung ưu tiên cách dùng thực tế cho quản trị hệ thống, HCNS, kế toán, Thu mua, trưởng bộ phận và nhân viên."
    )

    add_heading(doc, "2. Phân hệ Phòng ban", 1)
    add_heading(doc, "2.1. Mục đích", 2)
    add_para(
        doc,
        "Phân hệ Phòng ban dùng để quản lý cơ cấu tổ chức của công ty. Đây là nơi giữ bộ khung tổ chức, "
        "đầu mối quản lý nhân sự theo phòng, vai trò và quyền hạn của từng nhóm người dùng."
    )
    add_bullet(doc, "Tạo phòng ban gốc và phòng ban con.")
    add_bullet(doc, "Quản lý cây tổ chức của công ty.")
    add_bullet(doc, "Gán trưởng phòng hoặc người đứng đầu đơn vị.")
    add_bullet(doc, "Theo dõi nhân sự trong từng phòng.")
    add_bullet(doc, "Tạo vai trò và gán phân quyền theo từng phòng.")
    add_bullet(doc, "Điều chuyển nhân sự giữa các phòng ban.")

    add_heading(doc, "2.2. Các khu vực chính trên màn hình", 2)
    add_para(doc, "Màn hình Phòng ban hiện được tổ chức thành các khu vực chính sau:")
    add_table(
        doc,
        ["Khu vực", "Công dụng"],
        ROOM_SUMMARY_ROWS,
        [4.2, 12.3],
    )

    add_heading(doc, "2.3. Chức năng Tổng quan", 2)
    add_para(
        doc,
        "Khu vực Tổng quan giúp người dùng nắm nhanh bức tranh tổ chức toàn công ty mà không cần mở từng phòng."
    )
    add_bullet(doc, "Hiển thị tổng số phòng ban.")
    add_bullet(doc, "Hiển thị tổng số nhân sự toàn công ty.")
    add_bullet(doc, "Hiển thị số phòng đã có trưởng phòng và số phòng còn thiếu trưởng phòng.")
    add_bullet(doc, "Hiển thị phân bố giữa khối sản xuất và khối văn phòng.")
    add_bullet(doc, "Cho phép tìm theo tên hoặc mã phòng ban.")
    add_bullet(doc, "Có thể lọc nhanh các phòng thiếu trưởng phòng hoặc chưa có nhân sự.")
    add_bullet(doc, "Có thể đổi giữa chế độ danh sách và sơ đồ cây.")

    add_heading(doc, "2.4. Quản lý phòng ban", 2)
    add_para(doc, "Người có quyền phù hợp có thể thực hiện các nghiệp vụ chính sau:")
    add_number(doc, "Tạo mới phòng ban gốc hoặc phòng ban con.")
    add_number(doc, "Sửa thông tin phòng ban.")
    add_number(doc, "Xóa phòng ban khi không còn phù hợp.")
    add_number(doc, "Đổi cấp trên trong cây tổ chức.")
    add_number(doc, "Chỉ định trưởng phòng.")
    add_para(
        doc,
        "Nếu một phòng đã có người nhưng chưa có trưởng phòng, hệ thống hiển thị cảnh báo riêng để dễ xử lý. "
        "Phòng chưa có nhân sự được đánh dấu riêng, không bị coi là lỗi dữ liệu."
    )

    add_heading(doc, "2.5. Quản lý nhân sự trong phòng", 2)
    add_para(
        doc,
        "Tab Nhân sự trong Phòng ban là nơi xử lý nhân sự theo góc nhìn tổ chức, không phải góc nhìn lương."
    )
    add_bullet(doc, "Xem danh sách người thuộc phòng.")
    add_bullet(doc, "Lọc theo trạng thái nhân sự.")
    add_bullet(doc, "Chọn nhiều người để gán vai trò hàng loạt.")
    add_bullet(doc, "Chuyển nhiều người sang phòng khác.")
    add_bullet(doc, "Thêm nhân viên mới ngay trong màn hình Phòng ban.")
    add_para(
        doc,
        "Khi chuyển phòng, vai trò cũ có thể bị gỡ vì vai trò luôn thuộc về đúng phòng ban. "
        "Nếu người được chuyển đang là trưởng phòng cũ thì hệ thống cũng bỏ gán trưởng phòng ở đơn vị cũ."
    )

    add_heading(doc, "2.6. Tab Vai trò & Quyền", 2)
    add_para(
        doc,
        "Vai trò trong hệ thống là bó quyền gắn với từng phòng ban cụ thể. Một vai trò không dùng chung cho toàn công ty."
    )
    add_bullet(doc, "Tạo vai trò mới trong phòng.")
    add_bullet(doc, "Đổi tên vai trò.")
    add_bullet(doc, "Xóa vai trò.")
    add_bullet(doc, "Gán ma trận phân quyền cho vai trò.")
    add_para(
        doc,
        "Mỗi tài khoản chỉ giữ một vai trò tại một thời điểm, và vai trò đó phải thuộc đúng phòng của người dùng."
    )

    add_heading(doc, "2.7. Logic phân quyền", 2)
    add_para(doc, "Ma trận phân quyền hiện vận hành theo 3 lớp chính:")
    add_bullet(doc, "Xem: được vào module và đọc dữ liệu.")
    add_bullet(doc, "Thao tác: được thêm, sửa, xóa.")
    add_bullet(doc, "Phạm vi: quy định được đụng tới dữ liệu của ai.")
    add_para(doc, "Các mức phạm vi hiện có như sau:")
    add_table(doc, ["Phạm vi", "Ý nghĩa"], SCOPE_ROWS, [4.2, 12.3])
    add_para(
        doc,
        "Nguyên tắc đang áp dụng là: nếu đã có quyền thao tác thì hệ thống tự hiểu phải có quyền xem; "
        "nếu tắt quyền xem thì các quyền thao tác liên quan cũng bị tắt theo. Mục đích là tránh cấp quyền nửa chừng."
    )

    add_heading(doc, "2.8. Các quyền chi tiết quan trọng", 2)
    add_para(
        doc,
        "Ngoài CRUD cơ bản, phân hệ hiện dùng nhiều quyền chi tiết để kiểm soát các thao tác nhạy cảm."
    )
    add_bullet(doc, "Đặt trưởng phòng.")
    add_bullet(doc, "Đổi cấp trên trong cây tổ chức.")
    add_bullet(doc, "Gán vai trò.")
    add_bullet(doc, "Chuyển phòng ban.")
    add_bullet(doc, "Sửa ma trận phân quyền.")
    add_bullet(doc, "Khóa hoặc mở khóa tài khoản.")
    add_bullet(doc, "Đặt lại mật khẩu và thu hồi phiên đăng nhập.")
    add_bullet(doc, "Xem lương & BHXH khi được cấp quyền nhạy cảm.")

    add_heading(doc, "2.9. Những điều cần nhớ ở phân hệ Phòng ban", 2)
    add_bullet(doc, "Một tài khoản chỉ thuộc một phòng ban tại một thời điểm.")
    add_bullet(doc, "Một tài khoản chỉ giữ một vai trò tại một thời điểm.")
    add_bullet(doc, "Vai trò phải thuộc đúng phòng của người đó.")
    add_bullet(doc, "Chuyển phòng xong nên kiểm tra lại vai trò.")
    add_bullet(doc, "Phòng ban là nơi quản lý cơ cấu, không phải nơi khai chính sách lương.")

    add_heading(doc, "3. Phân hệ Lương", 1)
    add_heading(doc, "3.1. Mục đích", 2)
    add_para(
        doc,
        "Phân hệ Lương là trung tâm xử lý lương thời gian của hệ thống. Phân hệ này quản lý dữ liệu lương, "
        "cấu hình tính lương, bảng lương tháng, tạm ứng và phiếu lương cá nhân."
    )

    add_heading(doc, "3.2. Các tab chính", 2)
    add_table(doc, ["Tab", "Công dụng"], PAYROLL_TABS, [4.2, 12.3])

    add_heading(doc, "3.3. Tab Bảng lương tháng", 2)
    add_para(
        doc,
        "Đây là nơi người có quyền quản lý lương vận hành kỳ lương theo từng tháng."
    )
    add_number(doc, "Chọn tháng cần xử lý.")
    add_number(doc, "Sinh bảng lương từ dữ liệu chấm công, hồ sơ lương và cấu hình lương.")
    add_number(doc, "Soát từng dòng lương theo nhân viên.")
    add_number(doc, "Sửa các ô tay khi kỳ còn ở trạng thái nháp.")
    add_number(doc, "Chốt kỳ lương.")
    add_number(doc, "Đánh dấu đã chi trả.")
    add_number(doc, "Xuất bảng lương Excel hoặc file chuyển khoản ngân hàng.")
    add_number(doc, "In phiếu lương từng người.")

    add_heading(doc, "3.4. Trạng thái kỳ lương", 2)
    add_table(doc, ["Tên hiển thị", "Mã trạng thái", "Ý nghĩa"], PAYROLL_PERIOD_STATUS, [3.5, 3.2, 9.8])
    add_para(
        doc,
        "Chỉ kỳ lương ở trạng thái nháp mới được tính lại và sửa trực tiếp. "
        "Kỳ đã chốt hoặc đã chi được giữ lại để bảo toàn lịch sử."
    )

    add_heading(doc, "3.5. Tab Lương nhân viên", 2)
    add_para(
        doc,
        "Tab này giữ dữ liệu gốc của từng nhân viên để phục vụ tính lương. Đây là nơi khai mức lương và "
        "các khoản cố định theo từng người, không dùng kiểu cả tổ một mức tiền."
    )
    add_bullet(doc, "Lương cơ bản dùng làm mức đóng bảo hiểm.")
    add_bullet(doc, "Lương trách nhiệm.")
    add_bullet(doc, "Thưởng chuyên cần.")
    add_bullet(doc, "Phụ cấp thâm niên.")
    add_bullet(doc, "Phụ cấp khác.")
    add_bullet(doc, "Lương trả 1 lần (đợt 1).")
    add_bullet(doc, "Các khoản cộng hoặc trừ riêng theo từng người.")
    add_bullet(doc, "Cách tính thuế TNCN khi hồ sơ yêu cầu.")
    add_para(
        doc,
        "Các thay đổi tại đây sẽ ảnh hưởng tới những kỳ lương còn nháp. Những kỳ đã chốt hoặc đã chi sẽ không tự thay đổi ngược."
    )

    add_heading(doc, "3.6. Tab Cấu hình lương", 2)
    add_para(
        doc,
        "Đây là khu vực dữ liệu nhạy cảm, không phải ai vào module Lương cũng được xem hoặc sửa."
    )
    add_bullet(doc, "Quản lý tham số chung như giờ chuẩn, tỷ lệ thử việc, tăng ca.")
    add_bullet(doc, "Quản lý cấu hình theo bộ phận.")
    add_bullet(doc, "Quản lý tỷ lệ bảo hiểm người lao động và doanh nghiệp.")
    add_bullet(doc, "Quản lý giảm trừ gia cảnh, thuế và các ngưỡng liên quan.")
    add_para(
        doc,
        "Người có quyền xem cấu hình lương được đọc dữ liệu nhạy cảm; người có quyền cập nhật mới được lưu thay đổi."
    )

    add_heading(doc, "3.7. Tab Tạm ứng", 2)
    add_para(
        doc,
        "Tab Tạm ứng quản lý cả phiếu tạm ứng thông thường và phiếu lương đợt 1."
    )
    add_bullet(doc, "Tạm ứng: khoản ứng trước ngoài lương chính.")
    add_bullet(doc, "Lương đợt 1: khoản trả trước một phần lương tháng.")
    add_para(doc, "Các trạng thái của phiếu tạm ứng hiện có như sau:")
    add_table(doc, ["Tên hiển thị", "Mã trạng thái", "Ý nghĩa"], ADVANCE_STATUS, [3.5, 3.2, 9.8])
    add_para(
        doc,
        "Khi phiếu đã duyệt, hệ thống dùng số đó để trừ vào bảng lương. "
        "Lương đợt 1 không phải khoản thu nhập cộng thêm, mà là khoản trả trước sẽ khấu trừ lại."
    )

    add_heading(doc, "3.8. Phiếu lương của tôi", 2)
    add_para(
        doc,
        "Đây là khu vực self-service cho nhân viên, dùng để xem phiếu lương cá nhân mà không cần quyền quản trị."
    )
    add_bullet(doc, "Xem lương theo công.")
    add_bullet(doc, "Xem các khoản cộng.")
    add_bullet(doc, "Xem các khoản khấu trừ.")
    add_bullet(doc, "Xem số tạm ứng đã nhận.")
    add_bullet(doc, "Xem số thực lĩnh.")

    add_heading(doc, "3.9. Tạm ứng của tôi", 2)
    add_para(
        doc,
        "Khu vực này cho phép nhân viên tự tạo và theo dõi đề nghị tạm ứng của chính mình."
    )
    add_bullet(doc, "Tạo đề nghị tạm ứng.")
    add_bullet(doc, "Tạo đề nghị lương đợt 1 nếu doanh nghiệp áp dụng.")
    add_bullet(doc, "Theo dõi phiếu đang chờ duyệt, đã duyệt, bị từ chối hoặc đã hủy.")

    add_heading(doc, "3.10. Logic dữ liệu người dùng cần hiểu", 2)
    add_number(doc, "Bảng lương tháng lấy dữ liệu từ hồ sơ lương, chấm công và cấu hình lương.")
    add_number(doc, "Tạm ứng đã duyệt sẽ tự đi vào phần khấu trừ.")
    add_number(doc, "Kỳ nháp có thể tính lại và sửa tay.")
    add_number(doc, "Kỳ đã chốt hoặc đã chi phải giữ nguyên để bảo toàn lịch sử.")
    add_number(doc, "Phân hệ có cả số liệu tự tính và số liệu sửa tay, nên phải xem đúng trạng thái trước khi thao tác.")

    add_heading(doc, "4. Phân hệ Thu mua", 1)
    add_heading(doc, "4.1. Mục đích", 2)
    add_para(
        doc,
        "Phân hệ Thu mua dùng để tiếp nhận nhu cầu mua hàng từ các phòng ban, lập phiếu mua hàng gửi duyệt, "
        "theo dõi quá trình đã mua và xác nhận hàng đã nhận."
    )
    add_para(
        doc,
        "Luồng hiện tại được tách rõ: phòng ban tạo YCMH, bộ phận Thu mua lập PMH, người có quyền duyệt quyết định "
        "duyệt hay từ chối PMH, sau đó mới đi tiếp sang các bước tài chính ở phân hệ kế toán."
    )

    add_heading(doc, "4.2. Các khu vực chính", 2)
    add_table(doc, ["Khu vực", "Công dụng"], PURCHASE_AREAS, [4.2, 12.3])

    add_heading(doc, "4.3. Luồng nghiệp vụ chính", 2)
    add_number(doc, "Phòng ban tạo Yêu cầu mua hàng (YCMH) với ngày cần hàng, mục đích và danh sách vật tư.")
    add_number(doc, "Thu mua vào màn Mua hàng để xem danh sách YCMH chờ xử lý.")
    add_number(doc, "Mỗi Phiếu mua hàng (PMH) hiện tạo từ một YCMH nguồn để dễ truy vết.")
    add_number(doc, "Thu mua chọn nhà cung cấp phù hợp dựa trên dữ liệu mặt hàng và đơn giá đã quản lý.")
    add_number(doc, "Thu mua gửi PMH lên người có quyền duyệt.")
    add_number(doc, "Sau khi PMH được duyệt, Thu mua theo dõi phiếu đã mua và sau đó xác nhận đã nhận hàng.")
    add_number(doc, "Nếu hàng nhận thiếu, hệ thống vẫn ghi số nhận thực tế để làm căn cứ công nợ và thanh toán về sau.")

    add_heading(doc, "4.4. Trạng thái phiếu mua hàng", 2)
    add_table(doc, ["Tên hiển thị", "Mã trạng thái", "Ý nghĩa"], PURCHASE_STATUS, [3.5, 3.2, 9.8])
    add_para(
        doc,
        "Điểm quan trọng là YCMH và PMH là hai lớp chứng từ khác nhau: YCMH phản ánh nhu cầu từ phòng ban, "
        "PMH phản ánh quyết định mua thực tế của bộ phận Thu mua với nhà cung cấp."
    )

    add_heading(doc, "4.5. Nhà cung cấp và mặt hàng", 2)
    add_para(
        doc,
        "Danh mục Nhà cung cấp không chỉ lưu thông tin liên hệ mà còn lưu các mặt hàng, đơn vị tính, đơn giá và VAT "
        "mà từng nhà cung cấp đang bán."
    )
    add_bullet(doc, "Thu mua dùng dữ liệu này để so sánh nhanh giá giữa các nhà cung cấp cho cùng một vật tư.")
    add_bullet(doc, "Khi tạo PMH, hệ thống có thể gợi ý nhà cung cấp phù hợp theo mặt hàng đã quản lý.")
    add_bullet(doc, "Việc thương lượng thực tế với nhà cung cấp vẫn có thể diễn ra ngoài hệ thống; hệ thống chủ yếu ghi nhận kết quả cuối cùng.")

    add_heading(doc, "4.6. Những điều cần nhớ ở phân hệ Thu mua", 2)
    add_bullet(doc, "YCMH là nhu cầu phát sinh; PMH là phiếu mua thực tế do Thu mua lập.")
    add_bullet(doc, "Người lập PMH không nhất thiết là người duyệt PMH.")
    add_bullet(doc, "Đã nhận hàng là mốc quan trọng vì nó làm căn cứ cho đối chiếu công nợ và thanh toán phía sau.")
    add_bullet(doc, "Nên quản lý tốt danh mục mặt hàng theo nhà cung cấp để việc chọn nhà cung cấp nhanh và ít sai hơn.")

    add_heading(doc, "5. Diễn giải quyền chi tiết", 1)
    add_para(
        doc,
        "Phần này diễn giải theo đúng cách người dùng cuối hay hỏi: bật một quyền lên thì màn nào hiện ra, "
        "nút nào xuất hiện, và quyền đó phục vụ việc gì trong vận hành hằng ngày."
    )
    add_heading(doc, "5.1. Cách đọc ma trận quyền", 2)
    add_table(
        doc,
        ["Thành phần", "Giải thích thực tế"],
        PERMISSION_MATRIX_OVERVIEW,
        [4.2, 12.3],
    )
    add_para(
        doc,
        "Ví dụ: nếu một vai trò được bật quyền Xem của module Lương nhưng chưa bật Thao tác, người đó vẫn mở được "
        "màn Lương để xem bảng lương và tạm ứng, nhưng sẽ không sửa được dòng lương, không tính lại, không chốt kỳ."
    )

    add_heading(doc, "5.2. Quyền chi tiết của nhóm Phòng ban", 2)
    add_permission_table(
        doc,
        "5.2.1. Module Phòng ban",
        "Đây là module quản trị cơ cấu tổ chức. Người dùng thường gặp nhất ở đây là Admin và HCNS.",
        PHONG_BAN_PERMISSION_ROWS,
    )
    add_permission_table(
        doc,
        "5.2.2. Module Vai trò",
        "Module này nằm trong luồng Phòng ban vì vai trò luôn gắn với từng phòng cụ thể.",
        VAI_TRO_PERMISSION_ROWS,
    )
    add_permission_table(
        doc,
        "5.2.3. Module Người dùng",
        "Các quyền này thường được dùng ngay trong tab Tài khoản & Quyền của hồ sơ nhân sự hoặc khi điều phối nhân sự giữa các phòng.",
        NGUOI_DUNG_PERMISSION_ROWS,
    )

    add_heading(doc, "5.3. Quyền chi tiết của nhóm Nhân sự", 2)
    add_permission_table(
        doc,
        "5.3.1. Module Nhân sự",
        "Đây là nhóm quyền quan trọng nhất cho HCNS vì nó quyết định ai được xem và ai được sửa hồ sơ nhân sự tới mức nào.",
        NHAN_SU_PERMISSION_ROWS,
    )

    add_heading(doc, "5.4. Quyền chi tiết của nhóm Lương", 2)
    add_permission_table(
        doc,
        "5.4.1. Module Lương",
        "Nhóm quyền này điều khiển việc nhìn bảng lương, vận hành lương, duyệt tạm ứng, chốt kỳ và xuất file chi lương.",
        LUONG_PERMISSION_ROWS,
    )
    add_para(
        doc,
        "Điểm cần nhớ là phân hệ Lương đã tách quyền khá rõ: quyền xem không đồng nghĩa với quyền sửa, "
        "quyền duyệt tạm ứng không đồng nghĩa với quyền chốt kỳ, và quyền xuất file cũng tách riêng để tránh phát tán dữ liệu."
    )

    add_heading(doc, "5.5. Quyền chi tiết của nhóm Thu mua", 2)
    add_permission_table(
        doc,
        "5.5.1. Module Thu mua",
        "Nhóm quyền này điều khiển việc nhìn YCMH - PMH, lập phiếu mua hàng, duyệt hoặc hủy PMH.",
        THU_MUA_PERMISSION_ROWS,
    )
    add_para(
        doc,
        "Ở phân hệ Thu mua, cần phân biệt rõ người lập phiếu với người duyệt phiếu. Quyền Duyệt / từ chối PMH nên cấp rất chọn lọc để tránh tự lập rồi tự duyệt."
    )

    add_heading(doc, "6. Vai trò thường gặp", 1)
    add_para(
        doc,
        "Dưới đây là cách hiểu ngắn gọn theo nhóm người dùng phổ biến trong doanh nghiệp."
    )
    add_table(doc, ["Vai trò", "Mô tả thực tế"], ROLE_ROWS, [4.2, 12.3])

    add_heading(doc, "7. Các lưu ý vận hành quan trọng", 1)
    add_bullet(doc, "Không sửa dữ liệu lương nếu chưa xác định kỳ đó đang là nháp, đã chốt hay đã chi.")
    add_bullet(doc, "Không xem Phòng ban là nơi khai chính sách tiền lương; đó là việc của phân hệ Lương.")
    add_bullet(doc, "Sau khi chuyển phòng phải kiểm tra lại vai trò, vì vai trò cũ có thể không còn hợp lệ.")
    add_bullet(doc, "Dữ liệu cấu hình lương là dữ liệu nhạy cảm, không nên cấp tràn lan.")
    add_bullet(doc, "Phiếu lương đợt 1 và tạm ứng đều là khoản sẽ được trừ lại khi tính lương tháng.")
    add_bullet(doc, "Ở Thu mua cần quản lý chặt người nào được lập PMH, người nào được duyệt PMH và người nào chỉ được theo dõi.")
    add_bullet(doc, "Danh mục nhà cung cấp nên được cập nhật đúng mặt hàng và đơn giá để tránh chọn nhầm nhà cung cấp khi lập phiếu.")

    add_heading(doc, "8. Kết luận", 1)
    add_para(
        doc,
        "Phân hệ Phòng ban giúp doanh nghiệp quản lý cơ cấu, đầu mối chịu trách nhiệm, vai trò và phạm vi quyền hạn. "
        "Phân hệ Lương giúp doanh nghiệp quản lý tiền lương, tham số tính lương, bảng lương tháng và phiếu lương cá nhân. "
        "Phân hệ Thu mua giúp doanh nghiệp ghi nhận nhu cầu mua, lập phiếu mua thực tế, theo dõi duyệt mua và xác nhận nhận hàng."
    )
    add_para(
        doc,
        "Khi người dùng hiểu rõ ai thuộc phòng nào, ai giữ vai trò gì và ai được phép thao tác đến đâu, "
        "việc vận hành nhân sự, lương và thu mua sẽ rõ ràng, dễ truy vết và ít sai sót hơn."
    )

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    path = build_doc()
    print(path)
